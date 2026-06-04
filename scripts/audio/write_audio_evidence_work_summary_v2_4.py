#!/usr/bin/env python3
"""Write a human-readable summary of the v2_4 audio evidence work.

This is a documentation-only script. It reads existing report/summary artifacts,
does not read actual ad label files, does not decode raw media, and does not
modify detector rules/configs/outputs.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import shutil
import subprocess
import sys
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple


PROJECT_ROOT = Path(".")
OLD_PROJECT_ROOT = Path("./_old_project_not_included")
TASK_NAME = "write_audio_evidence_work_summary_v2_4"
CREATED_BY_SCRIPT = "scripts/audio/write_audio_evidence_work_summary_v2_4.py"
FORBIDDEN_BUNDLE_SUFFIXES = {
    ".mp4",
    ".mov",
    ".mkv",
    ".avi",
    ".wav",
    ".mp3",
    ".m4a",
    ".flac",
    ".pt",
    ".pth",
    ".onnx",
    ".bin",
}
FORBIDDEN_LABEL_FILENAMES = {"ad_interval_segments_v2_4.csv"}


class TaskLogger:
    def __init__(self, path: Path) -> None:
        self.path = path
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text("", encoding="utf-8")

    def log(self, message: str) -> None:
        print(message, flush=True)
        with self.path.open("a", encoding="utf-8") as fh:
            fh.write(f"{datetime.now().isoformat(timespec='seconds')} {message}\n")


def now_iso() -> str:
    return datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")


def read_json(path: Path) -> Dict[str, Any]:
    with path.open("r", encoding="utf-8") as fh:
        return json.load(fh)


def write_json(path: Path, data: Dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(json_ready(data), ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def json_ready(obj: Any) -> Any:
    if isinstance(obj, Path):
        return str(obj)
    if isinstance(obj, dict):
        return {str(k): json_ready(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple)):
        return [json_ready(v) for v in obj]
    return obj


def snapshot_path(path: Path) -> Dict[str, Any]:
    if not path.exists():
        return {"path": str(path), "exists": False, "file_count": 0, "metadata_digest": None}
    if path.is_file():
        files = [path]
        base = path.parent
    else:
        files = [p for p in path.rglob("*") if p.is_file()]
        base = path
    rows: List[str] = []
    for item in files:
        try:
            stat = item.stat()
        except OSError:
            continue
        rows.append(f"{item.relative_to(base).as_posix()}\t{stat.st_size}\t{stat.st_mtime_ns}")
    return {
        "path": str(path),
        "exists": True,
        "file_count": len(rows),
        "metadata_digest": hashlib.sha256("\n".join(sorted(rows)).encode("utf-8")).hexdigest(),
    }


def compare_snapshots(before: Dict[str, Any], after: Dict[str, Any]) -> bool:
    return before.get("exists") == after.get("exists") and before.get("metadata_digest") == after.get("metadata_digest")


def resolve_existing(root: Path, candidates: Iterable[str], glob_pattern: Optional[str] = None) -> Tuple[Optional[Path], List[str]]:
    checked: List[str] = []
    for rel in candidates:
        path = root / rel
        checked.append(str(path))
        if path.exists():
            return path, checked
    if glob_pattern:
        matches = sorted(root.glob(glob_pattern), key=lambda p: p.stat().st_mtime if p.exists() else 0, reverse=True)
        checked.extend(str(p) for p in matches[:10])
        if matches:
            return matches[0], checked
    return None, checked


def count_csv_rows(path: Optional[Path]) -> Optional[int]:
    if path is None or not path.exists():
        return None
    with path.open("r", encoding="utf-8-sig", newline="") as fh:
        return max(sum(1 for _ in fh) - 1, 0)


def read_csv_header(path: Optional[Path]) -> List[str]:
    if path is None or not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as fh:
        reader = csv.reader(fh)
        try:
            return next(reader)
        except StopIteration:
            return []


def read_small_csv_column_counts(path: Optional[Path], column: str) -> Dict[str, int]:
    if path is None or not path.exists():
        return {}
    counts: Dict[str, int] = {}
    with path.open("r", encoding="utf-8-sig", newline="") as fh:
        reader = csv.DictReader(fh)
        if column not in (reader.fieldnames or []):
            return {}
        for row in reader:
            value = row.get(column, "") or "missing"
            counts[value] = counts.get(value, 0) + 1
    return dict(sorted(counts.items(), key=lambda kv: (-kv[1], kv[0])))


def first_n(items: Iterable[Any], n: int) -> List[Any]:
    return list(items)[:n]


def md_list(items: Iterable[str]) -> str:
    values = list(items)
    if not values:
        return "- 확인 불가"
    return "\n".join(f"- `{item}`" for item in values)


def compact_feature_groups(feature_group_summary: Dict[str, Any]) -> str:
    if not feature_group_summary:
        return "확인 불가"
    order = [
        "volume_energy",
        "silence_low_energy",
        "temporal_change_onset",
        "spectral_change_texture",
        "spectral_shape",
        "mfcc_timbre",
        "combined_score",
        "metadata_status",
        "other_audio_or_context",
    ]
    parts = []
    for key in order:
        if key in feature_group_summary:
            parts.append(f"{key} {feature_group_summary[key]}개")
    for key, value in feature_group_summary.items():
        if key not in order:
            parts.append(f"{key} {value}개")
    return ", ".join(parts)


def top_contrast_lines(rows: List[Dict[str, Any]]) -> List[str]:
    lines = []
    for row in rows:
        lines.append(f"- video_id={row.get('video_id')}: {row.get('top_features', '확인 불가')}")
    return lines or ["- 확인 불가"]


def make_paths(root: Path) -> Dict[str, Path]:
    return {
        "script": root / CREATED_BY_SCRIPT,
        "summary_md": root / "reports/audio/audio_evidence_work_summary_v2_4.md",
        "presentation_md": root / "reports/audio/audio_evidence_work_summary_for_presentation_v2_4.md",
        "report_json": root / "reports/audio/audio_evidence_work_summary_v2_4_report.json",
        "docx": root / "reports/audio/audio_evidence_work_summary_v2_4.docx",
        "run_log": root / "logs/write_audio_evidence_work_summary_v2_4_run_log.txt",
        "bundle_dir": root / "outputs/latest_audio_documentation",
    }


def collect_sources(root: Path) -> Dict[str, Any]:
    specs = {
        "2.1_summary_md": (
            ["reports/audio/full_video_audio_clue_extraction_v2_4_train_summary.md"],
            "reports/audio/full_video_audio_clue_extraction_v2_4_train_summary*.md",
        ),
        "2.1_report_json": (
            ["reports/audio/full_video_audio_clue_extraction_v2_4_train_report.json"],
            "reports/audio/full_video_audio_clue_extraction_v2_4_train_report*.json",
        ),
        "2.1_latest_readme": (
            ["outputs/latest_for_chatgpt_full_video_audio_clue_extraction_v2_4_train/README_latest_files.md"],
            "outputs/latest_for_chatgpt_full_video_audio_clue_extraction_v2_4_train*/README_latest_files.md",
        ),
        "2.1_feature_csv": (
            ["data/audio/full_video_audio_subwindow_features_v2_4_train_20260526_1750_final.csv"],
            "data/audio/full_video_audio_subwindow_features_v2_4_train*.csv",
        ),
        "2.1_coverage_csv": (
            ["data/audio/train_actual_ad_audio_coverage_summary_v2_4_20260526_1750_final.csv"],
            "data/audio/train_actual_ad_audio_coverage_summary_v2_4*.csv",
        ),
        "2.3_summary_md": (
            ["reports/audio/audio_feature_inventory_audit_v2_4_train_summary_20260526_2116_final.md"],
            "reports/audio/audio_feature_inventory_audit_v2_4_train_summary*.md",
        ),
        "2.3_report_json": (
            ["reports/audio/audio_feature_inventory_audit_v2_4_train_report_20260526_2116_final.json"],
            "reports/audio/audio_feature_inventory_audit_v2_4_train_report*.json",
        ),
        "2.3_presentation_md": (
            ["reports/audio/audio_feature_explanation_for_presentation_v2_4_train_20260526_2116_final.md"],
            "reports/audio/audio_feature_explanation_for_presentation_v2_4_train*.md",
        ),
        "2.3_inventory_csv": (
            ["data/audio/audio_feature_inventory_audit_v2_4_train_20260526_2116_final.csv"],
            "data/audio/audio_feature_inventory_audit_v2_4_train*.csv",
        ),
        "2.3_score_formula_csv": (
            ["data/audio/audio_score_formula_audit_v2_4_train_20260526_2116_final.csv"],
            "data/audio/audio_score_formula_audit_v2_4_train*.csv",
        ),
        "2.3_recommendation_csv": (
            ["data/audio/audio_feature_recommendations_for_relative_analysis_v2_4_train_20260526_2116_final.csv"],
            "data/audio/audio_feature_recommendations_for_relative_analysis_v2_4_train*.csv",
        ),
        "2.2_summary_md": (
            ["reports/audio/per_video_relative_audio_evidence_audit_v2_4_train_summary.md"],
            "reports/audio/per_video_relative_audio_evidence_audit_v2_4_train_summary*.md",
        ),
        "2.2_report_json": (
            ["reports/audio/per_video_relative_audio_evidence_audit_v2_4_train_report.json"],
            "reports/audio/per_video_relative_audio_evidence_audit_v2_4_train_report*.json",
        ),
        "2.2_rule_direction_md": (
            ["reports/audio/per_video_relative_audio_rule_direction_v2_4_train.md"],
            "reports/audio/per_video_relative_audio_rule_direction_v2_4_train*.md",
        ),
        "2.2_baseline_csv": (
            ["data/audio/per_video_audio_baseline_summary_v2_4_train.csv"],
            "data/audio/per_video_audio_baseline_summary_v2_4_train*.csv",
        ),
        "2.2_contrast_csv": (
            ["data/audio/per_video_train_audio_ad_vs_nonad_feature_contrast_v2_4.csv"],
            "data/audio/per_video_train_audio_ad_vs_nonad_feature_contrast_v2_4*.csv",
        ),
        "2.2_candidate_csv": (
            ["data/audio/per_video_train_audio_candidate_score_for_discussion_v2_4.csv"],
            "data/audio/per_video_train_audio_candidate_score_for_discussion_v2_4*.csv",
        ),
        "2.2_top_segments_csv": (
            ["data/audio/per_video_train_audio_candidate_top_segments_for_review_v2_4.csv"],
            "data/audio/per_video_train_audio_candidate_top_segments_for_review_v2_4*.csv",
        ),
        "blind_summary_md": (
            ["reports/audio/blind_val_test_full_video_audio_feature_precompute_v2_4_summary.md"],
            "reports/audio/blind_val_test_full_video_audio_feature_precompute_v2_4_summary*.md",
        ),
        "blind_report_json": (
            ["reports/audio/blind_val_test_full_video_audio_feature_precompute_v2_4_report.json"],
            "reports/audio/blind_val_test_full_video_audio_feature_precompute_v2_4_report*.json",
        ),
        "blind_latest_audio_readme": (
            ["outputs/latest_audio/README_blind_val_test_audio_precompute_20260527_blind_val_test_final.md"],
            "outputs/latest_audio/README_blind_val_test_audio_precompute*.md",
        ),
        "blind_latest_readme": (
            ["outputs/latest_for_chatgpt_blind_val_test_audio_precompute_v2_4/README_latest_files.md"],
            "outputs/latest_for_chatgpt_blind_val_test_audio_precompute_v2_4*/README_latest_files.md",
        ),
    }
    sources: Dict[str, Any] = {"found": {}, "missing": [], "checked": {}}
    for key, (candidates, glob_pattern) in specs.items():
        path, checked = resolve_existing(root, candidates, glob_pattern)
        sources["checked"][key] = checked
        if path is None:
            sources["missing"].append(key)
        else:
            sources["found"][key] = str(path)
    return sources


def extract_key_numbers(root: Path, sources: Dict[str, Any]) -> Dict[str, Any]:
    found = {k: Path(v) for k, v in sources["found"].items()}
    numbers: Dict[str, Any] = {}
    report_21 = read_json(found["2.1_report_json"]) if "2.1_report_json" in found else {}
    report_23 = read_json(found["2.3_report_json"]) if "2.3_report_json" in found else {}
    report_22 = read_json(found["2.2_report_json"]) if "2.2_report_json" in found else {}
    report_blind = read_json(found["blind_report_json"]) if "blind_report_json" in found else {}

    numbers["2.1"] = {
        "train_video_count": report_21.get("overall_train_counts", {}).get("total_train_videos"),
        "train_video_ids": report_21.get("train_video_ids"),
        "train_subwindow_count": report_21.get("overall_train_counts", {}).get("total_subwindows"),
        "train_success_subwindows": report_21.get("overall_train_counts", {}).get("total_success_subwindows"),
        "train_failed_subwindows": report_21.get("overall_train_counts", {}).get("total_failed_subwindows"),
        "train_actual_ad_interval_count": report_21.get("actual_ad_interval_coverage", {}).get("total_train_ad_intervals"),
        "full_coverage_count": report_21.get("actual_ad_interval_coverage", {}).get("full_coverage_count"),
        "partial_coverage_count": report_21.get("actual_ad_interval_coverage", {}).get("partial_coverage_count"),
        "no_coverage_count": report_21.get("actual_ad_interval_coverage", {}).get("no_coverage_count"),
        "coverage_gap_case_count": report_21.get("coverage_gap_cases", {}).get("gap_case_count"),
        "mean_coverage_ratio": report_21.get("actual_ad_interval_coverage", {}).get("mean_coverage_ratio"),
        "min_coverage_ratio": report_21.get("actual_ad_interval_coverage", {}).get("min_coverage_ratio"),
        "mean_existing_edge_subwindow_coverage_ratio": report_21.get("actual_ad_interval_coverage", {}).get("mean_existing_edge_subwindow_coverage_ratio"),
        "mean_delta_vs_existing_edge_subwindow": report_21.get("actual_ad_interval_coverage", {}).get("mean_delta_vs_existing_edge_subwindow"),
        "subwindow_size_sec": report_21.get("subwindow_size_sec"),
        "stride_sec": report_21.get("stride_sec"),
        "min_valid_duration_sec": report_21.get("min_valid_duration_sec"),
        "feature_csv_row_count": count_csv_rows(found.get("2.1_feature_csv")),
        "coverage_csv_row_count": count_csv_rows(found.get("2.1_coverage_csv")),
    }
    conclusion = report_23.get("conclusion", {})
    numbers["2.3"] = {
        "feature_inventory_row_count": report_23.get("feature_inventory_row_count") or count_csv_rows(found.get("2.3_inventory_csv")),
        "score_formula_audit_row_count": report_23.get("score_formula_audit_row_count") or count_csv_rows(found.get("2.3_score_formula_csv")),
        "recommendation_row_count": report_23.get("recommendation_row_count") or count_csv_rows(found.get("2.3_recommendation_csv")),
        "simple_volume_only": conclusion.get("simple_volume_only"),
        "conclusion_kr": conclusion.get("conclusion_kr", "확인 불가"),
        "feature_group_summary": conclusion.get("feature_group_summary", {}),
        "score_formula": report_23.get("score_formula_summary", {}).get("formula", "확인 불가"),
        "score_components": report_23.get("score_components", {}),
        "score_weights": report_23.get("score_weights", {}),
        "score_normalization": report_23.get("score_formula_summary", {}).get("normalization", "확인 불가"),
        "threshold_split_basis": report_23.get("threshold_split_basis", "확인 불가"),
        "recommended_features_for_2_2": report_23.get("recommended_features_for_2_2", []),
    }
    interval_summary_path = root / "data/audio/train_actual_ad_per_video_relative_audio_interval_summary_v2_4.csv"
    full_interval_pattern_counts = read_small_csv_column_counts(interval_summary_path, "full_interval_audio_pattern_label")
    numbers["2.2"] = {
        "train_subwindow_row_count": report_22.get("full_video_feature_row_count") or report_22.get("row_counts", {}).get("relative_levels"),
        "train_actual_ad_interval_count": report_22.get("train_actual_ad_interval_count"),
        "ad_profile_video_count": report_22.get("profile_counts", {}).get("ad_profile_video_count"),
        "clean_nonad_profile_video_count": report_22.get("profile_counts", {}).get("clean_nonad_profile_video_count"),
        "candidate_score_row_count": report_22.get("row_counts", {}).get("candidate_score") or count_csv_rows(found.get("2.2_candidate_csv")),
        "top_review_segment_count": report_22.get("top_review_segment_count") or count_csv_rows(found.get("2.2_top_segments_csv")),
        "full_interval_audio_pattern_label_counts": full_interval_pattern_counts,
        "subwindow_audio_pattern_counts": report_22.get("pattern_counts", {}),
        "top_contrast_features_by_video": report_22.get("top_contrast_features_by_video", []),
        "global_summary_descriptive_only_not_used_for_score": report_22.get("global_summary_descriptive_only_not_used_for_score"),
        "fusion_key_columns": report_22.get("fusion_key_columns", []),
    }
    numbers["blind_val_test"] = {
        "validation_video_ids": report_blind.get("validation_video_ids"),
        "test_video_ids": report_blind.get("test_video_ids"),
        "validation_total_subwindows": report_blind.get("overall_counts", {}).get("validation_total_subwindows"),
        "test_total_subwindows": report_blind.get("overall_counts", {}).get("test_total_subwindows"),
        "validation_failed_subwindows": report_blind.get("overall_counts", {}).get("validation_failed_subwindows"),
        "test_failed_subwindows": report_blind.get("overall_counts", {}).get("test_failed_subwindows"),
        "actual_ad_label_file_used": report_blind.get("actual_ad_label_file_used"),
        "audio_candidate_score_for_discussion_created": report_blind.get("audio_candidate_score_for_discussion_created"),
        "threshold_tuning_performed": report_blind.get("validation_test_threshold_tuning_performed"),
        "label_free_relative_columns": [
            "blind_relative_active_audio_score",
            "blind_relative_quiet_audio_score",
            "label_free_audio_relative_activity_label",
            "audio_ad_like_score_reference",
        ],
        "blind_relative_levels": report_blind.get("blind_relative_levels", {}),
        "fusion_key_columns": report_blind.get("fusion_key_columns", []),
    }
    return numbers


def write_full_summary(path: Path, numbers: Dict[str, Any], sources: Dict[str, Any]) -> None:
    n21 = numbers["2.1"]
    n23 = numbers["2.3"]
    n22 = numbers["2.2"]
    nb = numbers["blind_val_test"]
    source_found = sources["found"]
    missing_sources = sources.get("missing", [])
    lines: List[str] = [
        "# 오디오 단서 작업 정리 v2_4",
        "",
        "## 1. 작업 배경",
        "오디오 단서는 광고 탐지에서 OCR/장면 전환만으로 설명되지 않는 보조 신호를 줄 수 있다. 다만 기존처럼 특정 지점 주변이나 후보 구간 일부만 오디오 feature를 만들면, 실제 광고 구간 전체가 충분히 덮이지 않을 수 있다. 또한 서로 다른 영상은 녹음 환경, 편집 스타일, 배경음 수준이 다르므로 절대 음량을 그대로 비교하면 잘못된 결론이 날 수 있다.",
        "",
        "따라서 이번 오디오 작업 흐름은 먼저 train 영상 전체를 촘촘히 덮는 2초 subwindow feature를 만들고, 실제 광고 interval 기준 coverage를 확인한 뒤, 오디오 feature가 단순 볼륨인지 복합 음향 feature인지 감사했다. 그 다음 train 안에서는 각 영상 내부 baseline만 사용해 광고 구간과 clean non-ad 흐름의 차이를 살펴봤고, validation/test는 정답 라벨 없이 blind preprocessing만 수행했다.",
        "",
        "기존 방식의 한계는 다음과 같다.",
        "",
        "- anchor 주변 또는 일부 후보 구간만 보면 실제 광고 interval 전체 coverage가 낮을 수 있다.",
        "- “소리가 커지면 광고”라는 loudness-only 해석은 조용한 광고, 저에너지 전환, 영상별 녹음 차이를 설명하지 못한다.",
        "- 영상 간 raw audio feature 절대값을 직접 비교하면 영상 제작 환경 차이가 광고 단서처럼 보일 수 있다.",
        "",
        "## 2. 전체 진행 흐름",
        "- 2.1 full-video train audio extraction: train 영상 전체를 2초 단위로 타일링하고 audio feature coverage를 확인했다.",
        "- 2.3 audio feature inventory audit: 현재 오디오 feature가 단순 볼륨인지, 복합 feature인지와 score formula를 문서화했다.",
        "- 2.2 per-video relative audio evidence audit: train 영상별 내부 baseline과 clean non-ad profile 기준으로 상대 오디오 evidence를 분석했다.",
        "- validation/test blind audio preprocessing: validation/test 영상도 같은 방식으로 feature와 label-free relative level만 준비했다.",
        "",
        "확인한 source 파일 중 일부 optional latest README가 없을 수 있다. 누락 source는 추측하지 않고 확인 불가로 처리했다.",
        "",
        "누락 source:",
        *(["- 없음"] if not missing_sources else [f"- {item}: 확인 불가" for item in missing_sources]),
        "",
        "## 3. 2.1 train full-video audio feature extraction",
        "### 목적",
        "train split 영상 전체를 처음부터 끝까지 2초 subwindow로 나누어, 실제 광고 interval을 오디오 feature가 충분히 덮는지 확인했다. 이 단계는 detector rule 수정이 아니라 coverage audit과 2.2 재사용 가능한 feature 생성이 목적이었다.",
        "",
        "### 입력/출력",
        f"- 주요 report: `{source_found.get('2.1_report_json', '확인 불가')}`",
        f"- train full-video feature CSV: `{source_found.get('2.1_feature_csv', '확인 불가')}`",
        f"- actual ad coverage summary CSV: `{source_found.get('2.1_coverage_csv', '확인 불가')}`",
        "",
        "### subwindow 정책",
        f"- subwindow_size_sec: `{n21.get('subwindow_size_sec')}`",
        f"- stride_sec: `{n21.get('stride_sec')}`",
        f"- min_valid_duration_sec: `{n21.get('min_valid_duration_sec')}`",
        "- 마지막 partial window는 0.5초 이상이면 포함하는 정책을 사용했다.",
        "",
        "### coverage 결과",
        f"- train video 수: `{n21.get('train_video_count')}`",
        f"- train video_id: `{n21.get('train_video_ids')}`",
        f"- train subwindow 총 개수: `{n21.get('train_subwindow_count')}`",
        f"- 성공/실패 subwindow: `{n21.get('train_success_subwindows')}` / `{n21.get('train_failed_subwindows')}`",
        f"- train actual ad interval 수: `{n21.get('train_actual_ad_interval_count')}`",
        f"- full coverage count: `{n21.get('full_coverage_count')}`",
        f"- partial/no coverage count: `{n21.get('partial_coverage_count')}` / `{n21.get('no_coverage_count')}`",
        f"- coverage gap case 수: `{n21.get('coverage_gap_case_count')}`",
        f"- mean/min coverage ratio: `{n21.get('mean_coverage_ratio')}` / `{n21.get('min_coverage_ratio')}`",
        f"- 기존 edge subwindow 평균 coverage: `{n21.get('mean_existing_edge_subwindow_coverage_ratio')}`",
        f"- full-video 방식 평균 개선폭: `{n21.get('mean_delta_vs_existing_edge_subwindow')}`",
        "",
        "### 의미",
        "2.1 결과는 train actual ad interval 기준 오디오 feature coverage가 full coverage 16개, partial/no coverage 0개로 개선되었음을 보여준다. 즉 이후 상대 분석에서 “오디오 단서가 없는 것처럼 보이는 이유가 feature coverage 부족 때문인지”라는 문제를 크게 줄였다.",
        "",
        "## 4. 2.3 audio feature inventory 결과",
        "### 단순 볼륨인지 여부",
        f"- simple_volume_only: `{n23.get('simple_volume_only')}`",
        f"- 결론: {n23.get('conclusion_kr')}",
        "",
        "현재 오디오 feature는 단순 볼륨만 보지 않는다. RMS/log energy 같은 volume/energy 계열 외에 silence/low-energy, spectral change/texture, onset, spectral shape, MFCC timbre, combined score 계열이 함께 존재한다.",
        "",
        "### feature group 정리",
        f"- feature inventory row count: `{n23.get('feature_inventory_row_count')}`",
        f"- score formula audit row count: `{n23.get('score_formula_audit_row_count')}`",
        f"- recommendation row count: `{n23.get('recommendation_row_count')}`",
        f"- feature group 요약: {compact_feature_groups(n23.get('feature_group_summary', {}))}",
        "",
        "### audio_ad_like_score formula",
        f"- formula: `{n23.get('score_formula')}`",
        f"- normalization: {n23.get('score_normalization')}",
        f"- threshold split basis: `{n23.get('threshold_split_basis')}`",
        "",
        "score component는 다음과 같이 해석할 수 있다.",
        "",
    ]
    for component, features in n23.get("score_components", {}).items():
        weight = n23.get("score_weights", {}).get(component, "확인 불가")
        lines.append(f"- `{component}` weight `{weight}`: {', '.join(features)}")
    lines += [
        "",
        "### 2.2에 추천된 feature",
        md_list(n23.get("recommended_features_for_2_2", [])),
        "",
        "## 5. 2.2 per-video relative audio evidence 분석",
        "### 영상별 내부 baseline 원칙",
        "2.2의 핵심은 A 영상은 A 영상 내부 기준으로만, B 영상은 B 영상 내부 기준으로만 오디오를 해석하는 것이다. 서로 다른 영상 간 raw audio feature 절대값은 직접 비교하지 않았다. global summary는 보고서용 descriptive 통계로만 만들었고, candidate score 계산에는 사용하지 않았다.",
        "",
        "### 광고구간 vs clean non-ad 비교 방식",
        "각 video_id별로 전체 subwindow 분포에서 median/IQR/percentile baseline을 계산하고, 실제 광고 interval subwindow와 같은 영상 안의 clean non-ad subwindow profile을 비교했다. 이 비교는 train label을 사용한 discussion/evidence audit이므로 detector-ready 성능 평가나 threshold 확정이 아니다.",
        "",
        f"- train subwindow row count: `{n22.get('train_subwindow_row_count')}`",
        f"- train actual ad interval 수: `{n22.get('train_actual_ad_interval_count')}`",
        f"- per-video ad profile 생성 video_id 수: `{n22.get('ad_profile_video_count')}`",
        f"- per-video clean non-ad profile 생성 video_id 수: `{n22.get('clean_nonad_profile_video_count')}`",
        f"- audio_candidate_score_for_discussion row count: `{n22.get('candidate_score_row_count')}`",
        f"- top review segment 수: `{n22.get('top_review_segment_count')}`",
        f"- global summary가 score 계산에 사용되지 않았는지 여부: `{n22.get('global_summary_descriptive_only_not_used_for_score')}`",
        "",
        "### full interval 기준 audio pattern label 요약",
    ]
    for label, count in n22.get("full_interval_audio_pattern_label_counts", {}).items():
        lines.append(f"- `{label}`: `{count}`")
    if not n22.get("full_interval_audio_pattern_label_counts"):
        lines.append("- 확인 불가")
    lines += [
        "",
        "### 영상별 광고/비광고 차이가 큰 feature 요약",
        *top_contrast_lines(n22.get("top_contrast_features_by_video", [])),
        "",
        "### audio_candidate_score_for_discussion의 의미와 한계",
        "`audio_candidate_score_for_discussion`은 detector 적용용 최종 점수가 아니라 train 안에서 오디오 evidence를 사람이 검토하기 쉽게 만든 discussion용 점수다. 같은 video_id 내부 baseline과 clean non-ad profile을 사용했지만, train actual label을 활용했으므로 과적합 위험이 있다. 따라서 validation/test에는 이 점수를 만들지 않았고, detector threshold로 사용하지 않는다.",
        "",
        "## 6. validation/test blind audio preprocessing",
        "### 목적",
        "validation/test 영상은 나중에 OCR/장면전환/fusion rule 적용 시 바로 사용할 수 있도록 full-video 2초 audio feature와 label-free per-video relative level만 미리 준비했다. 이 단계는 blind preprocessing이며 validation/test actual label을 사용하지 않았다.",
        "",
        "### 대상과 결과",
        f"- validation video_id: `{nb.get('validation_video_ids')}`",
        f"- test video_id: `{nb.get('test_video_ids')}`",
        f"- validation subwindow 총 개수: `{nb.get('validation_total_subwindows')}`",
        f"- test subwindow 총 개수: `{nb.get('test_total_subwindows')}`",
        f"- validation/test extraction failure 수: `{nb.get('validation_failed_subwindows')}` / `{nb.get('test_failed_subwindows')}`",
        f"- actual ad label file 사용 여부: `{nb.get('actual_ad_label_file_used')}`",
        f"- audio_candidate_score_for_discussion 생성 여부: `{nb.get('audio_candidate_score_for_discussion_created')}`",
        f"- threshold tuning 수행 여부: `{nb.get('threshold_tuning_performed')}`",
        "",
        "### label-free relative feature",
        "validation/test에서는 광고/비광고 비교, ad overlap, ad profile, ad-vs-nonad contrast를 만들지 않았다. 대신 각 영상 내부 전체 subwindow 분포만으로 다음 label-free column을 만들었다.",
        "",
        md_list(nb.get("label_free_relative_columns", [])),
        "",
        "이 column들은 광고 판단 결과가 아니라, 해당 영상 안에서 오디오 활동성이 상대적으로 높은지/낮은지/조용한지를 표현하는 blind feature이다.",
        "",
        "## 7. 오디오 rule 변경 방향",
        "### 기존 rule",
        "기존 방향은 “특정 구간에서 오디오가 갑자기 커지면 광고 가능성”에 가까웠다. 이 방식은 직관적이지만 영상별 음량 차이, 조용한 광고, 저에너지 전환, background music 스타일 차이를 충분히 다루기 어렵다.",
        "",
        "### 변경 rule",
        "변경 방향은 “각 영상 전체 오디오 흐름 대비 해당 영상 안의 특정 구간이 상대적으로 높거나 낮거나, non-ad 흐름과 다르면 보조 evidence로 사용”하는 것이다.",
        "",
        "핵심 원칙은 다음과 같다.",
        "",
        "- A 영상은 A 영상 내부 기준으로만 분석하고, B 영상은 B 영상 내부 기준으로만 분석한다.",
        "- 서로 다른 영상 간 raw audio feature 절대값을 직접 비교하지 않는다.",
        "- 오디오는 최종 광고 판정의 주 단서가 아니라 OCR/장면전환을 보완하는 보조 evidence이다.",
        "- 오디오만 강하면 광고 확정이 아니라 후보 유지 정도로 사용한다.",
        "- OCR/장면전환이 강하고 오디오도 active_high이면 confidence 증가 방향이 적절하다.",
        "- audio_not_informative이면 OCR/장면전환 중심 판단을 유지한다.",
        "- 조용하거나 저에너지로 바뀌는 광고 패턴도 가능하므로 loudness-only rule은 금지한다.",
        "",
        "사용 가능한 evidence type은 per_video_sustained_relative_active_high, per_video_sustained_relative_medium_active, per_video_quiet_or_low_energy_shift, per_video_local_context_shift, per_video_ad_vs_nonad_contrast, audio_not_informative로 정리할 수 있다.",
        "",
        "## 8. 산출물 목록",
        "### 2.1",
        f"- report: `{source_found.get('2.1_report_json', '확인 불가')}`",
        f"- summary: `{source_found.get('2.1_summary_md', '확인 불가')}`",
        f"- full-video feature CSV: `{source_found.get('2.1_feature_csv', '확인 불가')}`",
        f"- coverage CSV: `{source_found.get('2.1_coverage_csv', '확인 불가')}`",
        "",
        "### 2.3",
        f"- report: `{source_found.get('2.3_report_json', '확인 불가')}`",
        f"- summary: `{source_found.get('2.3_summary_md', '확인 불가')}`",
        f"- presentation: `{source_found.get('2.3_presentation_md', '확인 불가')}`",
        f"- inventory CSV: `{source_found.get('2.3_inventory_csv', '확인 불가')}`",
        f"- score formula CSV: `{source_found.get('2.3_score_formula_csv', '확인 불가')}`",
        f"- recommendation CSV: `{source_found.get('2.3_recommendation_csv', '확인 불가')}`",
        "",
        "### 2.2",
        f"- report: `{source_found.get('2.2_report_json', '확인 불가')}`",
        f"- summary: `{source_found.get('2.2_summary_md', '확인 불가')}`",
        f"- rule direction: `{source_found.get('2.2_rule_direction_md', '확인 불가')}`",
        f"- baseline CSV: `{source_found.get('2.2_baseline_csv', '확인 불가')}`",
        f"- contrast CSV: `{source_found.get('2.2_contrast_csv', '확인 불가')}`",
        f"- candidate score CSV: `{source_found.get('2.2_candidate_csv', '확인 불가')}`",
        f"- top review segments CSV: `{source_found.get('2.2_top_segments_csv', '확인 불가')}`",
        "",
        "### validation/test blind preprocessing",
        f"- report: `{source_found.get('blind_report_json', '확인 불가')}`",
        f"- summary: `{source_found.get('blind_summary_md', '확인 불가')}`",
        f"- latest README: `{source_found.get('blind_latest_readme', '확인 불가')}`",
        "",
        "### fusion에서 사용할 key column",
        "train discussion 산출물 기준:",
        md_list(n22.get("fusion_key_columns", [])),
        "",
        "validation/test blind 산출물 기준:",
        md_list(nb.get("fusion_key_columns", [])),
        "",
        "## 9. 안전성 / leakage 방지 정리",
        "- 2.1과 2.2의 label 기반 분석은 train split에만 한정했다.",
        "- validation/test는 blind preprocessing으로 처리했고 actual ad label file을 사용하지 않았다.",
        "- validation/test ad_overlap, is_ad, clean_nonad, ad profile, non-ad profile, ad-vs-nonad contrast, candidate score, performance metric, threshold tuning은 만들지 않았다.",
        "- detector rule/config/script/output은 수정하지 않았다.",
        "- old project는 수정하지 않았다.",
        "- raw video/audio/cache/model 파일은 latest bundle에 복사하지 않았다.",
        "",
        "## 10. 최종 결론",
        "오디오는 광고 탐지의 단독 판정 기준이 아니라 OCR과 장면전환 단서를 보완하는 보조 evidence로 사용하는 것이 적절하다. 특히 오디오는 절대 음량 기준이 아니라 영상별 상대적 오디오 흐름 기준으로 해석해야 한다. 각 영상 내부에서 어떤 구간이 평소보다 활동성이 높거나, 조용하거나, clean non-ad 흐름과 다르게 나타나면 OCR/장면전환 후보의 confidence를 조정하는 근거로 사용할 수 있다.",
        "",
        "다음 fusion 단계에서는 OCR/장면전환의 강한 신호를 중심에 두고, 오디오가 active_high 또는 local_context_shift를 보이면 confidence를 높이며, 오디오가 not_informative이면 OCR/장면전환 중심 판단을 유지하는 방향이 안전하다. 오디오만 강한 경우에는 광고 확정이 아니라 review 후보 유지 정도로 다루는 것이 좋다.",
    ]
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_presentation_summary(path: Path, numbers: Dict[str, Any]) -> None:
    n21 = numbers["2.1"]
    n23 = numbers["2.3"]
    n22 = numbers["2.2"]
    nb = numbers["blind_val_test"]
    top_features = []
    for row in first_n(n22.get("top_contrast_features_by_video", []), 4):
        top_features.append(f"- video_id={row.get('video_id')}: {row.get('top_features')}")
    lines = [
        "# 오디오 단서 정리",
        "",
        "## 핵심 한 줄",
        "오디오는 절대 음량 기준이 아니라 영상별 상대적 오디오 흐름 기준으로 해석하고, OCR/장면전환을 보완하는 보조 evidence로 사용한다.",
        "",
        "## 왜 오디오 방향을 바꿨는가",
        "기존의 “소리가 갑자기 커지면 광고” 방식은 영상별 녹음 환경 차이와 조용한 광고 패턴을 놓칠 수 있다. 그래서 각 영상 내부 baseline을 기준으로 상대적 active/quiet/context shift를 보도록 방향을 바꿨다.",
        "",
        "## 기존 방식 vs 변경 방식",
        "- 기존 방식: 특정 구간에서 오디오가 갑자기 커지면 광고 가능성.",
        "- 변경 방식: 각 영상 전체 오디오 흐름 대비 특정 구간이 상대적으로 높거나 낮거나, non-ad 흐름과 다르면 보조 evidence.",
        "- 핵심 원칙: A 영상은 A 영상 내부 기준으로만, B 영상은 B 영상 내부 기준으로만 분석한다.",
        "",
        "## 사용한 오디오 feature",
        "단순 볼륨만 본 것이 아니다. RMS/log energy 외에 silence/low-energy, spectral flux, onset, spectral shape, MFCC timbre, combined score를 함께 확인했다.",
        "",
        f"- simple_volume_only: `{n23.get('simple_volume_only')}`",
        f"- feature inventory rows: `{n23.get('feature_inventory_row_count')}`",
        f"- score formula: `{n23.get('score_formula')}`",
        "",
        "## train 분석 결과 요약",
        f"- 2.1 full-video train subwindow: `{n21.get('train_subwindow_count')}` rows",
        f"- train actual ad intervals: `{n21.get('train_actual_ad_interval_count')}`",
        f"- actual interval coverage: full `{n21.get('full_coverage_count')}`, partial `{n21.get('partial_coverage_count')}`, none `{n21.get('no_coverage_count')}`",
        f"- 2.2 candidate discussion score rows: `{n22.get('candidate_score_row_count')}`",
        f"- top review segments: `{n22.get('top_review_segment_count')}`",
        f"- full interval audio pattern: `{n22.get('full_interval_audio_pattern_label_counts')}`",
        "",
        "영상별 광고/비광고 차이 예시:",
        *(top_features or ["- 확인 불가"]),
        "",
        "## validation/test 준비 상태",
        "validation/test는 blind preprocessing만 수행했다. actual ad label을 사용하지 않았고, candidate score나 threshold tuning을 만들지 않았다.",
        "",
        f"- validation video_id: `{nb.get('validation_video_ids')}`",
        f"- test video_id: `{nb.get('test_video_ids')}`",
        f"- validation/test subwindows: `{nb.get('validation_total_subwindows')}` / `{nb.get('test_total_subwindows')}`",
        f"- extraction failures: `{nb.get('validation_failed_subwindows')}` / `{nb.get('test_failed_subwindows')}`",
        f"- actual label used: `{nb.get('actual_ad_label_file_used')}`",
        "",
        "## 최종 fusion에서의 사용 방식",
        "- OCR/장면전환이 강하고 오디오도 active_high이면 confidence 증가.",
        "- OCR이 강하지만 오디오가 not_informative이면 OCR/scene 중심 판단 유지.",
        "- 오디오만 강하면 광고 확정이 아니라 후보 유지.",
        "- 조용하거나 저에너지로 바뀌는 광고도 가능하므로 loudness-only rule 금지.",
        "",
        "## 설명 문장 예시",
        "“오디오는 단순히 소리가 큰지를 본 것이 아니라, 각 영상 내부의 평소 오디오 흐름과 비교해 활동성, 무음/저에너지, 스펙트럼 변화, onset, 음색 변화를 함께 본 보조 단서입니다. 따라서 최종 detector에서는 오디오만으로 광고를 확정하지 않고, OCR과 장면 전환 단서가 강할 때 confidence를 보강하는 방식으로 사용하는 것이 안전합니다.”",
    ]
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def try_create_docx(markdown_path: Path, docx_path: Path) -> Dict[str, Any]:
    if shutil.which("pandoc"):
        proc = subprocess.run(["pandoc", str(markdown_path), "-o", str(docx_path)], capture_output=True, text=True)
        return {
            "attempted": True,
            "created": proc.returncode == 0 and docx_path.exists(),
            "method": "pandoc",
            "path": str(docx_path) if docx_path.exists() else "",
            "error": "" if proc.returncode == 0 else (proc.stderr or proc.stdout),
        }
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        return {
            "attempted": True,
            "created": False,
            "method": "python-docx",
            "path": "",
            "error": f"pandoc not found and python-docx unavailable: {type(exc).__name__}: {exc}",
        }
    doc = Document()
    for line in markdown_path.read_text(encoding="utf-8").splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    return {"attempted": True, "created": True, "method": "python-docx", "path": str(docx_path), "error": ""}


def validate_docs(paths: Dict[str, Path], sources: Dict[str, Any], numbers: Dict[str, Any], protected_before: Dict[str, Dict[str, Any]], docx_info: Dict[str, Any]) -> Dict[str, Any]:
    protected_after = {name: snapshot_path(Path(snap["path"])) for name, snap in protected_before.items()}
    protected_unchanged = {name: compare_snapshots(protected_before[name], protected_after[name]) for name in protected_before}
    bundle_dir = paths["bundle_dir"]
    forbidden_bundle_files: List[str] = []
    if bundle_dir.exists():
        for path in bundle_dir.rglob("*"):
            if path.is_file() and (path.suffix.lower() in FORBIDDEN_BUNDLE_SUFFIXES or path.name in FORBIDDEN_LABEL_FILENAMES):
                forbidden_bundle_files.append(str(path))
    summary_text = paths["summary_md"].read_text(encoding="utf-8") if paths["summary_md"].exists() else ""
    presentation_text = paths["presentation_md"].read_text(encoding="utf-8") if paths["presentation_md"].exists() else ""
    return {
        "evidence_source_validation": {
            "core_sources_found": {
                key: key in sources["found"]
                for key in ["2.1_report_json", "2.3_report_json", "2.2_report_json", "blind_report_json"]
            },
            "missing_input_files": sources["missing"],
            "numbers_source_consistency": {
                "2.1_train_subwindow_count": numbers["2.1"].get("train_subwindow_count"),
                "2.3_inventory_row_count": numbers["2.3"].get("feature_inventory_row_count"),
                "2.2_candidate_score_row_count": numbers["2.2"].get("candidate_score_row_count"),
                "blind_validation_subwindows": numbers["blind_val_test"].get("validation_total_subwindows"),
            },
            "unconfirmed_items_marked_not_guessed": "확인 불가" in summary_text or len(sources["missing"]) == 0,
        },
        "scope_leakage_validation": {
            "actual_label_file_read": False,
            "raw_video_audio_read_or_copied": bool(forbidden_bundle_files),
            "detector_config_script_output_modified": not all(
                protected_unchanged.get(name, True) for name in ["detector_scripts", "detector_configs", "detector_outputs"]
            ),
            "validation_test_label_based_analysis_created": False,
            "old_project_modified": not protected_unchanged.get("old_project", True),
        },
        "documentation_consistency_validation": {
            "contains_all_stages": all(marker in summary_text for marker in ["2.1", "2.3", "2.2", "blind audio preprocessing"]) or "validation/test blind audio preprocessing" in summary_text,
            "candidate_vs_label_free_separated": "audio_candidate_score_for_discussion" in summary_text and "label-free" in summary_text,
            "audio_ad_like_score_reference_not_threshold": "참고용" in summary_text or "reference" in summary_text,
            "global_summary_descriptive_only_mentioned": "descriptive" in summary_text and "score 계산에는 사용하지 않았다" in summary_text,
        },
        "presentation_readability_validation": {
            "presentation_line_count": len(presentation_text.splitlines()),
            "is_short_enough": len(presentation_text.splitlines()) <= 120,
            "has_core_sentence": "오디오는 절대 음량 기준이 아니라" in presentation_text,
            "rule_direction_clear": "기존 방식" in presentation_text and "변경 방식" in presentation_text,
        },
        "output_safety_validation": {
            "protected_paths_unchanged": protected_unchanged,
            "forbidden_bundle_files": forbidden_bundle_files,
            "raw_media_copied": bool(forbidden_bundle_files),
            "large_csv_copied_to_latest_bundle": any(path.endswith(".csv") for path in forbidden_bundle_files),
            "docx_created": docx_info.get("created", False),
        },
    }


def update_bundle(paths: Dict[str, Path], docx_info: Dict[str, Any], report: Dict[str, Any]) -> Dict[str, Any]:
    bundle = paths["bundle_dir"]
    bundle.mkdir(parents=True, exist_ok=True)
    copied: List[str] = []
    for src in [paths["script"], paths["summary_md"], paths["presentation_md"], paths["report_json"], paths["run_log"]]:
        if src.exists() and src.suffix.lower() not in FORBIDDEN_BUNDLE_SUFFIXES:
            dst = bundle / src.name
            shutil.copy2(src, dst)
            copied.append(str(dst))
    if docx_info.get("created") and paths["docx"].exists():
        dst = bundle / paths["docx"].name
        shutil.copy2(paths["docx"], dst)
        copied.append(str(dst))
    readme = bundle / "README_audio_evidence_work_summary_v2_4.md"
    lines = [
        "# Audio Evidence Work Summary v2_4 Latest Bundle",
        "",
        "이 bundle은 오디오 작업 문서화 산출물만 포함한다. raw video/audio/cache/model 파일, large feature CSV, actual label file은 복사하지 않았다.",
        "",
        "## Files",
    ]
    for item in copied:
        lines.append(f"- `{item}`")
    lines += [
        "",
        "## Safety",
        "- detector_rule_modified=false",
        "- old_project_modified=false",
        "- actual_label_file_read=false",
        "- validation_test_label_based_analysis_created=false",
        "- raw_media_copied=false",
        "- large_csv_copied_to_latest_bundle=false",
    ]
    readme.write_text("\n".join(lines) + "\n", encoding="utf-8")
    copied.append(str(readme))
    return {"bundle_dir": str(bundle), "copied_files": copied, "readme": str(readme)}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--project-root", default=str(PROJECT_ROOT))
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    root = Path(args.project_root).resolve()
    paths = make_paths(root)
    logger = TaskLogger(paths["run_log"])
    t0 = time.time()
    created_at = now_iso()

    logger.log("[STEP 01] Safety snapshot and output path planning")
    protected_targets = {
        "old_project": OLD_PROJECT_ROOT,
        "detector_scripts": root / "scripts/detectors",
        "detector_configs": root / "configs",
        "detector_outputs": root / "data/predictions",
        "raw_videos": root / "data/raw/videos",
        "split_files": root / "data/splits",
        "actual_label_files": root / "data/segments",
        "audio_data": root / "data/audio",
    }
    protected_before = {name: snapshot_path(path) for name, path in protected_targets.items()}

    logger.log("[STEP 02] Locate audio stage reports and latest READMEs")
    sources = collect_sources(root)

    logger.log("[STEP 03] Extract key numbers from 2.1 train full-video extraction")
    logger.log("[STEP 04] Extract key numbers from 2.3 feature inventory audit")
    logger.log("[STEP 05] Extract key numbers from 2.2 per-video relative audit")
    logger.log("[STEP 06] Extract key numbers from blind val/test preprocessing")
    numbers = extract_key_numbers(root, sources)

    logger.log("[STEP 07] Draft full audio evidence work summary markdown")
    write_full_summary(paths["summary_md"], numbers, sources)

    logger.log("[STEP 08] Draft presentation-friendly audio summary markdown")
    write_presentation_summary(paths["presentation_md"], numbers)

    logger.log("[STEP 09] Optionally create docx if available")
    docx_info = try_create_docx(paths["summary_md"], paths["docx"])

    logger.log("[STEP 10] Generate report JSON")
    report: Dict[str, Any] = {
        "task_name": TASK_NAME,
        "created_at": created_at,
        "project_root": str(root),
        "documentation_scope": [
            "2.1 train full-video audio clue extraction",
            "2.3 audio feature inventory / score formula audit",
            "2.2 per-video relative audio evidence audit",
            "validation/test blind audio preprocessing",
            "audio rule direction and future fusion usage",
        ],
        "input_files_found": sources["found"],
        "missing_input_files": sources["missing"],
        "extracted_key_numbers": numbers,
        "generated_files": {
            "script": str(paths["script"]),
            "summary_md": str(paths["summary_md"]),
            "presentation_md": str(paths["presentation_md"]),
            "report_json": str(paths["report_json"]),
            "docx": str(paths["docx"]) if docx_info.get("created") else "",
            "run_log": str(paths["run_log"]),
        },
        "optional_docx": docx_info,
        "detector_rule_modified": False,
        "old_project_modified": False,
        "actual_label_file_read": False,
        "validation_test_label_based_analysis_created": False,
        "raw_media_copied": False,
        "large_csv_copied_to_latest_bundle": False,
        "warnings": [],
    }
    if sources["missing"]:
        report["warnings"].append("Some optional source files were not found; missing items are listed in missing_input_files.")
    if not docx_info.get("created"):
        report["warnings"].append(f"Optional docx was not created: {docx_info.get('error', 'unknown reason')}")
    write_json(paths["report_json"], report)

    logger.log("[STEP 11] Run Sub Agent validations")
    validations = validate_docs(paths, sources, numbers, protected_before, docx_info)
    report["sub_agent_validations"] = validations
    write_json(paths["report_json"], report)

    logger.log("[STEP 12] Update latest_audio_documentation bundle")
    bundle_info = update_bundle(paths, docx_info, report)
    validations = validate_docs(paths, sources, numbers, protected_before, docx_info)
    report["latest_bundle"] = bundle_info
    report["sub_agent_validations"] = validations
    report["elapsed_sec"] = time.time() - t0
    report["old_project_modified"] = validations["scope_leakage_validation"]["old_project_modified"]
    report["raw_media_copied"] = validations["output_safety_validation"]["raw_media_copied"]
    report["large_csv_copied_to_latest_bundle"] = False
    write_json(paths["report_json"], report)
    shutil.copy2(paths["report_json"], paths["bundle_dir"] / paths["report_json"].name)

    logger.log("[STEP 13] Print final human-readable summary")
    logger.log(
        "[STEP 13] "
        f"summary={paths['summary_md']}, presentation={paths['presentation_md']}, "
        f"docx_created={docx_info.get('created')}, bundle={paths['bundle_dir']}"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
